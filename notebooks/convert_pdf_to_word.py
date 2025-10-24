#!/usr/bin/env python3
"""
ocr_label_pdf_to_docx.py

OCR + smart option splitting for MCQ exam PDFs.

Usage examples:
    # process pages 1-5 test chunk
    python ocr_label_pdf_to_docx.py --input "Exam All-Questions.pdf" --outdir output --start 1 --end 5

    # process entire file in chunks of 20 pages and merge them
    python ocr_label_pdf_to_docx.py --input "Exam All-Questions.pdf" --outdir output --chunk 20 --merge

Author: ChatGPT (script provided for local execution)
"""

import os, sys, io, re, math, argparse, textwrap
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import fitz  # PyMuPDF
from tqdm import tqdm

# try import pytesseract, provide friendly error
try:
    import pytesseract
except Exception as e:
    pytesseract = None

# ---------------------------
# Helper & parsing functions
# ---------------------------

def ensure_tesseract():
    if pytesseract is None:
        raise RuntimeError("pytesseract is not installed. pip install pytesseract and ensure Tesseract OCR binary is installed.")
    # optionally allow user to set tesseract_cmd externally if needed
    return True

def render_page_image(page, zoom=2.0):
    """Render PyMuPDF page to PIL Image at given zoom"""
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    return img

def crop_left(image, left_fraction):
    """Crop left fraction (0..1) from PIL Image"""
    w, h = image.size
    crop_x = int(w * left_fraction)
    return image.crop((crop_x, 0, w, h))

def ocr_image(image, lang="eng", psm=None):
    """Run OCR on PIL image using pytesseract. psm = pagesegmode if needed."""
    if pytesseract is None:
        return ""
    config = ""
    if psm:
        config = f'--psm {psm}'
    text = pytesseract.image_to_string(image, lang=lang, config=config)
    return text

# Heuristics for parsing questions+options
_qnum_re = re.compile(r'^\s*([0-9]{1,3})\.\s*(.*)')  # lines that start like "1. ..."

# Detect option tokens e.g. "A.", "(a)", "1)", "a)" or line likely starting an option
_opt_token_re = re.compile(r'^\s*(?:\(?([A-Da-d0-9])\)?[\.\)\:\-]|\b[A-Da-d]\b[\.\)]?)\s*(.*)')

def split_options_by_heuristics(text_after_question):
    """
    Given a string following a question mark, attempt to split into 3-6 option candidates.
    Heuristics:
      - If explicit tokens like 'A.' '(a)' '1)' are present, split on those.
      - Else, split on patterns of double spaces or capitalized verbs/significant words when chunk length is moderate.
      - Fallback: split by common separators when there are multiple segments.
    """
    lines = [ln.strip() for ln in re.split(r'\r\n|\n|\r', text_after_question) if ln.strip()]
    joined = " ".join(lines)

    # 1) try explicit markers A., (a), 1), etc.
    tokens = re.split(r'(?:(?<=\s)|^)(?:\(?[A-Da-d]\)?[\.\)\:\-]|[0-9]{1,2}\))\s+', joined)
    # the split above will produce empty first element if starts with token, so filter empties
    tokens = [t.strip() for t in tokens if t.strip()]
    if len(tokens) >= 2:
        return tokens

    # 2) Look for sequences that look like options separated by two or more spaces (common OCR result)
    if re.search(r'\s{2,}', joined):
        parts = [p.strip() for p in re.split(r'\s{2,}', joined) if p.strip()]
        if 2 <= len(parts) <= 6:
            return parts

    # 3) Split where a capitalized word starts and previous segment length is short
    # (heuristic: split on " Bring", " Convert", " The", " It", etc.)
    split_points = re.split(r'(?:(?<=\.)\s+|(?<=\?)\s+|(?<=:)\s+)', joined)
    # If that gave many pieces, try grouping into options of reasonable length
    candidates = [p.strip() for p in re.split(r'(?<=\bBring\b|(?<=\bBring\b)|(?<=\bConvert\b)|(?<=\bMake\b)|(?<=\bSelect\b)|(?<=\bThe\b)|(?<=\bIf\b))', joined) if p.strip()]
    # Fallback to splitting by " Bring" or " The"
    if len(candidates) >= 2 and len(candidates) <= 6:
        cleaned = [c.strip(" -:;,.") for c in candidates if c.strip()]
        if 2 <= len(cleaned) <= 6:
            return cleaned

    # 4) If long but contains repeats of 'or' maybe delimited by ' or '
    if joined.count(" or ") >= 2:
        parts = [p.strip() for p in joined.split(" or ") if p.strip()]
        if len(parts) >= 2:
            return parts

    # Final fallback: try to chop into 4 roughly-equal segments by sentence punctuation
    sents = re.split(r'\.\s+', joined)
    if len(sents) >= 2 and len(sents) <= 6:
        return [s.strip() for s in sents if s.strip()]

    # If nothing works, return empty list to indicate no detected options
    return []

def label_options(opts):
    """Return list of labeled options (a..z) given list of strings"""
    labeled = []
    for i, o in enumerate(opts):
        label = f"({chr(ord('a') + i)})"
        labeled.append((label, o.strip()))
    return labeled

def clean_text_lines(lines, cleanup_phrases=None):
    """Remove lines containing any of the cleanup phrases (case-insensitive) and strip whitespace"""
    if cleanup_phrases is None:
        cleanup_phrases = ["Clear Ans", "Previous", "Next", "min left", "min remaining", "Time left", "Answer Options", "Course", "Select any", "Select the"]
    out = []
    for ln in lines:
        ok = True
        for ph in cleanup_phrases:
            if ph.lower() in ln.lower():
                ok = False
                break
        if ok and ln.strip():
            out.append(ln.strip())
    return out

# ---------------------------
# Document building functions
# ---------------------------

def write_questions_to_docx(document, page_num, questions, page_image_bytes=None):
    """Append page image (if provided) and question blocks to python-docx Document"""
    para = document.add_paragraph(f'--- Page {page_num} ---')
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    if page_image_bytes:
        document.add_picture(io.BytesIO(page_image_bytes), width=Inches(6))
    document.add_paragraph("")  # blank

    for q in questions:
        qnum, qtext, opts = q
        # question header
        header = f"{qnum}." if qnum else ""
        p = document.add_paragraph()
        run = p.add_run(f"{header} {qtext}".strip())
        run.bold = True
        document.add_paragraph("")  # blank between question and options
        if opts:
            labeled = label_options(opts)
            for lab, opttext in labeled:
                p_opt = document.add_paragraph()
                p_opt.add_run(f"{lab} {opttext}")
        else:
            document.add_paragraph("(No explicit options detected on this question.)")
        document.add_paragraph("")  # blank line between questions

def merge_docx_files(file_list, merged_output_path):
    """Naive merge of docx files by appending paragraphs from each into a new doc"""
    merged = Document()
    for fpath in file_list:
        doc = Document(fpath)
        # append a page break between files
        for element in doc.element.body:
            merged.element.body.append(element)
    merged.save(merged_output_path)

# ---------------------------
# Main processing per chunk
# ---------------------------

def process_chunk(pdf_path, out_dir, start_page=1, end_page=None, crop_left_fraction=0.18,
                  ocr_lang="eng", psm=3, dpi_zoom=2.0, cleanup_phrases=None):
    """
    Process pages from start_page..end_page (1-based inclusive).
    Saves a .docx file in out_dir named chunk_<start>_<end>.docx
    """
    ensure_tesseract()
    doc = fitz.open(pdf_path)
    total_pages = doc.page_count
    if end_page is None:
        end_page = total_pages
    start_idx = max(0, start_page - 1)
    end_idx = min(end_page, total_pages)  # end_page inclusive

    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    output_docx = out_dir / f"chunk_{start_page:04d}_{end_page:04d}.docx"
    document = Document()
    # Set default font size for readability
    style = document.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    document.add_heading(f'OCR Labeled: pages {start_page} to {end_page}', level=1)
    document.add_paragraph(f'Crop left fraction: {crop_left_fraction}. OCR lang: {ocr_lang}. PSM: {psm}.')

    for pnum in range(start_idx, end_idx):
        page = doc.load_page(pnum)
        # render image
        pil_img = render_page_image(page, zoom=dpi_zoom)
        cropped_img = crop_left(pil_img, crop_left_fraction) if crop_left_fraction and crop_left_fraction > 0 else pil_img
        # OCR
        ocr_text = ocr_image(cropped_img, lang=ocr_lang, psm=psm)
        if not ocr_text or len(ocr_text.strip()) < 10:
            # fallback to PDF-extract
            ocr_text = page.get_text("text")
        # basic cleanup
        lines = [ln for ln in re.split(r'\r\n|\n|\r', ocr_text) if ln.strip()]
        lines = clean_text_lines(lines, cleanup_phrases)
        joined = " ".join(lines).strip()
        # parse questions: look for qnum markers "12." OR for '?' marks
        questions = []
        # Approach A: scan lines sequentially for explicit "N." beginnings
        i = 0
        nlines = len(lines)
        while i < nlines:
            ln = lines[i]
            m = _qnum_re.match(ln)
            if m:
                qnum = m.group(1)
                qrest = m.group(2).strip()
                # collect following lines until next qnum or blank
                j = i + 1
                block = [qrest] if qrest else []
                while j < nlines and not _qnum_re.match(lines[j]):
                    block.append(lines[j])
                    j += 1
                i = j
                # Now block likely contains question + options; split by '?' if present
                block_text = " ".join(block).strip()
                # if a '?' exists inside qrest+block, split at first ? to separate question & options
                if '?' in block_text:
                    q_before, q_after = block_text.split('?', 1)
                    q_text = (m.group(2) + ' ' + q_before).strip() if m.group(2) else q_before.strip()
                    opts = split_options_by_heuristics(q_after.strip())
                else:
                    # no explicit '?', try splitting heuristically: assume last 3-5 short lines are options
                    # consider the last up to 6 lines and try splitting
                    maybe_opts = []
                    # find short lines towards end
                    toks = block
                    # if explicit markers in any token, use them
                    joined_block = " ".join(toks)
                    detected_opts = split_options_by_heuristics(joined_block)
                    if detected_opts:
                        # No qnum-specific question text easily isolated; set question = first portion or m.group(2)
                        q_text = m.group(2) if m.group(2) else " ".join(toks[:1])
                        opts = detected_opts
                    else:
                        # fallback: treat everything as question text
                        q_text = " ".join(toks)
                        opts = []
                questions.append((qnum, q_text, opts))
            else:
                # not start of qnum - try to find '?' inside the line
                if '?' in ln:
                    # take rest of following lines as candidate options
                    q_before, q_after = ln.split('?', 1)
                    q_text = q_before.strip() + '?'
                    opts = split_options_by_heuristics(q_after.strip())
                    # also attach subsequent short lines as candidate options if none found
                    j = i + 1
                    extra = []
                    while j < nlines and len(extra) < 10:
                        if len(lines[j].split()) <= 20:
                            extra.append(lines[j].strip())
                        j += 1
                    if not opts and extra:
                        opts = split_options_by_heuristics(" ".join(extra))
                    questions.append((None, q_text, opts))
                    i += 1
                else:
                    i += 1
                    continue

        # If nothing parsed, put the cleaned text as single block to allow manual correction later
        if not questions:
            # attempt to find sentence with question mark globally
            if '?' in joined:
                # split on '?' into separate question blocks
                parts = [p.strip() for p in joined.split('?') if p.strip()]
                for ix, p in enumerate(parts):
                    if ix == 0:
                        # first part treat as question text
                        # last part may be options in subsequent segment(s)
                        pass
            # fallback: one block
            questions = [(None, joined[:500], [])]  # keep first 500 chars for sanity

        # convert PIL image to bytes for docx insertion
        bio = io.BytesIO()
        cropped_img.save(bio, format="PNG")
        page_image_bytes = bio.getvalue()

        # write to docx
        write_questions_to_docx(document, pnum + 1, questions, page_image_bytes=page_image_bytes)

    # Save docx
    document.save(output_docx)
    return str(output_docx)

# ---------------------------
# Command-line interface
# ---------------------------

def parse_args():
    parser = argparse.ArgumentParser(description="OCR + auto-label MCQ options from PDF to DOCX")
    parser.add_argument("--input", "-i", required=True, help="Input PDF path")
    parser.add_argument("--outdir", "-o", default="ocr_output", help="Output directory")
    parser.add_argument("--start", type=int, default=1, help="Start page (1-based)")
    parser.add_argument("--end", type=int, default=None, help="End page (1-based inclusive). Default = end of file")
    parser.add_argument("--chunk", type=int, default=20, help="Pages per chunk when processing entire file")
    parser.add_argument("--crop", type=float, default=0.18, help="Fraction of page width to crop from left (0.0-0.4 recommended)")
    parser.add_argument("--lang", type=str, default="eng", help="Tesseract language (default 'eng')")
    parser.add_argument("--psm", type=int, default=3, help="Tesseract Page Segmentation Mode (3 = auto)")
    parser.add_argument("--zoom", type=float, default=2.0, help="Rendering zoom for page images (2.0 recommended)")
    parser.add_argument("--merge", action="store_true", help="If processing in chunks, merge chunk docx outputs into one merged.docx")
    return parser.parse_args()

def main():
    args = parse_args()
    pdf_path = args.input
    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    # Attempt to ensure Tesseract presence
    if pytesseract is None:
        print("Warning: pytesseract not installed. Install with `pip install pytesseract` and ensure Tesseract binary is installed.")
    else:
        # optional: print tesseract binary location
        try:
            print("Tesseract version:", pytesseract.get_tesseract_version())
        except Exception:
            pass

    # If user passed both start & end, just do single chunk
    if args.end is not None:
        print(f"Processing pages {args.start} to {args.end} ...")
        output_docx = process_chunk(pdf_path, outdir, start_page=args.start, end_page=args.end,
                                    crop_left_fraction=args.crop, ocr_lang=args.lang, psm=args.psm,
                                    dpi_zoom=args.zoom)
        print("Saved:", output_docx)
        return

    # Otherwise process entire PDF in chunks
    doc = fitz.open(pdf_path)
    total_pages = doc.page_count
    chunk_size = max(1, args.chunk)
    chunk_files = []
    for start in range(1, total_pages + 1, chunk_size):
        end = min(start + chunk_size - 1, total_pages)
        print(f"Processing chunk pages {start}-{end} ...")
        out_path = process_chunk(pdf_path, outdir, start_page=start, end_page=end,
                                crop_left_fraction=args.crop, ocr_lang=args.lang, psm=args.psm,
                                dpi_zoom=args.zoom)
        chunk_files.append(out_path)
        print(f"Chunk saved to {out_path}")
    if args.merge and len(chunk_files) > 1:
        merged = outdir / "merged_ocr_labeled.docx"
        print("Merging chunk files to", merged)
        merge_docx_files(chunk_files, merged)
        print("Merged file saved:", merged)

if __name__ == "__main__":
    main()
